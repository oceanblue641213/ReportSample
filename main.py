import replace_text as rt
from db_connection import create_connection as conn
import datetime
from dotenv import load_dotenv
from system_sub_categories import get_system_sub_categories_data
from site_modules import get_site_modules_data
from onbording_ratings import get_onbording_ratings_data
from company_assets import get_company_assets_data
from ghg_reports import get_ghg_reports_data
from new_total_emission import get_new_total_emission_data
import transfer_language as tl
from table_handler import (
    organization_boundaries_setting_range_2_1_1, 
    ghg_emission_list_2_3_1, 
    indirect_ghg_emission_evaluation_2_3_3,
    seven_ghg_emission_3_2_1,
    ghg_emission_list_3_2_2,
    direct_ghg_emission_3_3_1,
    indirect_ghg_emission_3_3_3
    )
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document

def main():
    # input_company_id = input('輸入公司id：')
    # input_year = input('輸入欲查詢年份：')
    company_id = "652cf86ce76f185628e48d9e" # 64a39d240db2a9a8f4a50893 # 63d0ac0a22d4cdf11191097e
    year = "2025"
    today = datetime.date.today()
    formatted_date = today.strftime("%Y年%m月%d日")
    
    load_dotenv()  # 讀取 .env 文件中的環境變數
    conn_string = os.getenv('DB_CONNECTIONSTRING')
    db_name = os.getenv('DB_NAME')
    redis_host = os.getenv('REDIS_HOST', 'localhost')
    redis_port = int(os.getenv('REDIS_PORT', 6379))
    redis_password = os.getenv('REDIS_PASSWORD', '12345678')
    
    with conn(conn_string, db_name) as db:
        try:
            all_collections = db.get_all_collections()
            site_modules, company_assets, onboarding_ratings, new_total_emission, ghg_report, system_sub_categories, new_translations = (
            all_collections['site_modules'],
            all_collections['company_assets'],
            all_collections['onboarding_ratings'],
            all_collections['new_total_emission'],
            all_collections['ghg_report'],
            all_collections['system_sub_categories'],
            all_collections['new_translations']
            )
            
            all_categories, categories_without_1 = get_system_sub_categories_data(system_sub_categories)
            all_site_modules_datas = get_site_modules_data(site_modules, company_id, fields=['companyName', 'levelStatus', 'usedRegion', 'address', 'EFGroup'])
            all_site_modules_ids = [item.get('_id') for item in all_site_modules_datas]
            sorted_data = sorted(all_site_modules_datas, key=lambda x: x.get('levelStatus') != 'parent')
            
            parent = next((item for item in all_site_modules_datas if item.get('levelStatus') == 'parent'), None)
            parent_rating = get_onbording_ratings_data(onboarding_ratings, parent.get('_id'))
            all_company_assets = get_company_assets_data(company_assets, all_site_modules_ids)
            all_ghg_reports_datas = get_ghg_reports_data(ghg_report, parent.get('_id'), fields=['emissionFactors']).get('emissionFactors')
            totals, ghg_emission_tables, direct_ghg_emission = get_new_total_emission_data(new_total_emission, all_categories, all_site_modules_ids)
            
            i18n_service = tl.I18nService(
                mongodb_uri=conn_string,
                db_name=db_name,
                collection_name='new_translations',
                redis_host='localhost',
                redis_port=redis_port,
                redis_password=redis_password
            )
            
            products = [{
                'Industry': 'Industry',
                'Carbon Footprint': 'Carbon Footprint'
            }]

            # 翻譯產品資料
            translated_products = i18n_service.translate_documents(products, "tw")
            
            print("開始:")
            
            # 計算
            totalStationaryCombustionValue = ghg_emission_tables['totalStationaryCombustionValue']
            totalIndustrialProcesses = ghg_emission_tables['totalIndustrialProcesses']
            totalMobileCombustionValue = ghg_emission_tables['totalMobileCombustionValue']
            totalRefrigerantsFugitives = ghg_emission_tables['totalRefrigerantsFugitives']
            totalCategories2Value = ghg_emission_tables['totalCategories2Value']
            totalCategories3Value = ghg_emission_tables['totalCategories3Value']
            totalCategories4Value = ghg_emission_tables['totalCategories4Value']
            category1TotalEmissionValue = totalStationaryCombustionValue + totalIndustrialProcesses + totalMobileCombustionValue + totalRefrigerantsFugitives
            totalEmissionValue = category1TotalEmissionValue + totalCategories2Value + totalCategories3Value + totalCategories4Value
            category1TotalEmissionRatio = round(category1TotalEmissionValue/totalEmissionValue, 2)
            
            category1TotalAmountEmission = direct_ghg_emission['totalGWPEmissionValue']
            largestAmountEmission = direct_ghg_emission['co2GWPEmissionValue']
            largestAmountEmissionRatio = round(largestAmountEmission/category1TotalAmountEmission, 2)
            secondLargestAmountEmission = direct_ghg_emission['ch4GWPEmissionValue']
            secondLargestAmountEmissionRatio = round(secondLargestAmountEmission/category1TotalAmountEmission, 2)
            
            replacements = {
            '[company]': parent.get('companyName'),
            '[year]': parent_rating.get('baseYear'),
            '[report date]': formatted_date,
            '[covered_range_from_to]': 
                parent_rating.get('baseYearPeriodStart').strftime('%Y年%m月%d日')
                + '至'
                + parent_rating.get('baseYearPeriodEnd').strftime('%Y年%m月%d日'),
            '[onbording_ratings_type]': parent_rating.get('type'),
            '[threshold]': str(parent_rating.get('threshold')),
            '[total_ghg_amount]': str(totals['totalGWPEmissionValue']),
            '[total_co2_amount]': str(totals['co2GWPEmissionValue']),
            '[total_co2_ration]': str(round(totals['co2GWPEmissionValue']/totals['totalGWPEmissionValue'], 2)),
            '[total_ch4_amount]': str(totals['ch4GWPEmissionValue']),
            '[total_ch4_ration]': str(round(totals['ch4GWPEmissionValue']/totals['totalGWPEmissionValue'], 2)),
            '[direct_ghg_emission]': str(category1TotalEmissionValue),
            '[direct_ghg_emission_ratio]': str(round(category1TotalEmissionValue/totalEmissionValue, 2)),
            '[largestAmountEmission]': str(largestAmountEmission),
            '[largestAmountEmissionRatio]': str(largestAmountEmissionRatio),
            '[secondLargestAmountEmission]': str(secondLargestAmountEmission),
            '[secondLargestAmountEmissionRatio]': str(secondLargestAmountEmissionRatio),
            '[category1TotalAmountEmission]': str(category1TotalAmountEmission),
            '[category1TotalEmissionRatio]': str(category1TotalEmissionRatio),
            '[largestAmountEmission]': str(largestAmountEmission),
            '[largestAmountEmissionRatio]': str(largestAmountEmissionRatio),
            '[secondLargestAmountEmission]': str(secondLargestAmountEmission),
            '[secondLargestAmountEmissionRatio]': str(secondLargestAmountEmissionRatio),
            }
            
            # 打開文檔
            doc = Document('溫室氣體盤查報告書v2.0.docx')
            
            # 替換文本
            rt.replace_text_in_word(doc, replacements)
            
            # 替換表格
            table00_data = organization_boundaries_setting_range_2_1_1(sorted_data)
            replace_table(doc.tables[0], table00_data)
            
            table01_data = ghg_emission_list_2_3_1(all_ghg_reports_datas)
            replace_table(doc.tables[1], table01_data, n = 2)
            
            table02_data = indirect_ghg_emission_evaluation_2_3_3(parent_rating, categories_without_1)
            replace_table(doc.tables[2], table02_data, n = 1, h = 2)
            
            table04_data = seven_ghg_emission_3_2_1(doc.tables[4], totals)
            
            table05_data = ghg_emission_list_3_2_2(doc.tables[5], ghg_emission_tables, category1TotalEmissionValue, totalEmissionValue)
            
            table06_data = direct_ghg_emission_3_3_1(all_company_assets)
            replace_table(doc.tables[6], table06_data)
            
            table07_data = seven_ghg_emission_3_2_1(doc.tables[7], direct_ghg_emission)
            
            table08_data = indirect_ghg_emission_3_3_3(all_company_assets)
            replace_table(doc.tables[8], table08_data)
            
            print("結束")
    
            # 保存新文檔
            doc.save('updated_document.docx')
            
        except Exception as e:
            print(f"發生錯誤：{e}")
    

def replace_table(table, new_data, n: int = 1, h: int = 0):
    """
    替換表格內容，保留標題行並根據數據動態調整表格大小
    
    Args:
        table: Word文檔中的表格對象
        new_data: 要插入的新數據（二維列表）
        n: 標題行的行數
        h: 要跳過的列數
    """
    # 獲取需要的行數（數據行數）
    required_rows = len(new_data) + n  # 因為要保留標題行
    current_rows = len(table.rows)
    
    # 調整表格大小
    if required_rows > current_rows:
        # 需要添加行
        for _ in range(required_rows - current_rows):
            table.add_row()
    elif required_rows < current_rows:
        # 需要刪除多餘的行
        for _ in range(current_rows - required_rows):
            table._element.remove(table.rows[-1]._element)
    
    # 然後填充數據行
    for i, row_data in enumerate(new_data, start=n):  # start=n 跳過標題行
        for j, cell_data in enumerate(row_data, start=h):
            if j < len(table.columns):  # 確保不超過列數
                cell = table.cell(i, j)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                cell_data = str(cell_data)
                if cell_data == "True":
                    cell_data = "V"
                elif cell_data == "False":
                    cell_data = ""
                else:
                    pass
                run = paragraph.add_run(cell_data)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


if __name__ == '__main__':
    main()


# def transfer_language(translations, new_translations):
#         print("開始翻譯")
#         # 讀取所有的語系資料
#         translations = translations.find()

#         # 迭代每個語系資料，轉換格式
#         for translation in translations:
#             for lang_data in translation['tran']:  # 每個語系資料
#                 lang = lang_data['name']  # 獲取語系名稱（如 "en", "de", 等）
#                 translations_data = lang_data['data']  # 獲取對應語系的翻譯資料
                
#                 # 構建新結構
#                 new_translation = {
#                     "lang": lang,
#                     "translations": translations_data
#                 }
                
#                 # 插入或更新資料到 new_translations 表格
#                 new_translations.update_one(
#                     {"lang": lang},  # 根據語系查詢
#                     {"$set": new_translation},  # 如果已經存在相同語系，則進行覆蓋
#                     upsert=True  # 如果找不到資料，則新增
#                 )
#         print("結束翻譯")