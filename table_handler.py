from docx.enum.text import WD_ALIGN_PARAGRAPH

def organization_boundaries_setting_range_2_1_1(data):
    result = [(item.get('companyName'), item.get('address', "")) for item in data]
    
    return result

def ghg_emission_list_2_3_1(data):
    result = [(item.get('category'), item.get('fuelType', ""), item.get('co2EmissionValue', ""), item.get('ch4EmissionValue', ""), item.get('n2oEmissionValue', ""), item.get('hfcsEmissionValue', ""), item.get('pfcsEmissionValue', ""), item.get('sf6EmissionValue', ""), item.get('nf3EmissionValue', "")) for item in data]
    
    return result

def indirect_ghg_emission_evaluation_2_3_3(data, categories):
    """
    如果遵從要求(reqByLaw)==true，後面欄位都不用填，最後一欄是否量化計算填"V"
    如果遵德要求(reqByLaw)==false，後面欄位都要填，最後一欄是否量化計算填"X"
    """
    ids = [str(item['_id']) for item in categories]
    query = [item for item in data.get('rating') if str(item.get('subCategory')) in ids]
    threshold = data.get('threshold')
    result = [
                (
                    "是" if item.get('reqByLaw') == True else "否",
                    "" if item.get('reqByLaw') == True else ("" if item.get('rating1') is None or item.get('rating1') == "" else item.get('rating1')),
                    "" if item.get('reqByLaw') == True else ("" if item.get('rating2') is None or item.get('rating2') == "" else item.get('rating2')),
                    "" if item.get('reqByLaw') == True else ("" if item.get('rating3') is None or item.get('rating3') == "" else item.get('rating3')),
                    "" if item.get('reqByLaw') == True else ("" if item.get('rating4') is None or item.get('rating4') == "" else item.get('rating4')),
                    "" if item.get('reqByLaw') == True else ("" if item.get('rating5') is None or item.get('rating5') == "" else item.get('rating5')),
                    "" if item.get('reqByLaw') == True else int(sum(float(item.get(f'rating{i}')) if item.get(f'rating{i}') is not None and item.get(f'rating{i}') != "" else 0 for i in range(1, 6))),
                    "V" if (item.get('reqByLaw') == True or sum(float(item.get(f'rating{i}')) if item.get(f'rating{i}') is not None and item.get(f'rating{i}') != "" else 0 for i in range(1, 6)) >= threshold) else "X",
                ) for item in query
            ]
    
    return result

def seven_ghg_emission_3_2_1(table, raw_data):
    total_gwp = raw_data['totalGWPEmissionValue']
    
    # 取出其他的值並計算每個值占總值的百分比
    values_and_percentages = [
        [value, str(round((value / total_gwp) * 100, 2)) + '%'] 
        for key, value in raw_data.items() 
    ]
    
    keys = [key for key in raw_data]
    values = [raw_data[key] for key in keys]
    percentages = [str(round((value / total_gwp) * 100, 2)) + '%' for value in values]
    data = [values, percentages]
    
    for i, row_data in enumerate(data, start=1):
        for j, cell_data in enumerate(row_data, start=1):
            if j < len(table.columns):  # 確保不超過列數
                cell = table.cell(i, j)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                cell_data = str(cell_data)
                run = paragraph.add_run(cell_data)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def ghg_emission_list_3_2_2(table, raw_data, category1_total, total_emissions):
    table.cell(2,1).text = ""
    table.cell(2,1).text = str(category1_total)
    table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2,5).text = ""
    table.cell(2,5).text = str(raw_data['totalCategories2Value'])
    table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2,6).text = ""
    table.cell(2,6).text = str(raw_data['totalCategories3Value'])
    table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2,7).text = ""
    table.cell(2,7).text = str(raw_data['totalCategories4Value'])
    table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2,8).text = ""
    table.cell(2,8).text = str(total_emissions)
    table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4,1).text = ""
    table.cell(4,1).text = str(f"{(category1_total/total_emissions*100):.2f}%")
    table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4,5).text = ""
    table.cell(4,5).text = str(f"{(raw_data['totalCategories2Value']/total_emissions*100):.2f}%")
    table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4,6).text = ""
    table.cell(4,6).text = str(f"{(raw_data['totalCategories3Value']/total_emissions*100):.2f}%")
    table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4,7).text = ""
    table.cell(4,7).text = str(f"{(raw_data['totalCategories4Value']/total_emissions*100):.2f}%")
    table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5,1).text = ""
    table.cell(5,1).text = str(f"{(raw_data['totalStationaryCombustionValue']/total_emissions*100):.2f}%")
    table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5,2).text = ""
    table.cell(5,2).text = str(f"{(raw_data['totalIndustrialProcesses']/total_emissions*100):.2f}%")
    table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5,3).text = ""
    table.cell(5,3).text = str(f"{(raw_data['totalMobileCombustionValue']/total_emissions*100):.2f}%")
    table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5,4).text = ""
    table.cell(5,4).text = str(f"{(raw_data['totalRefrigerantsFugitives']/total_emissions*100):.2f}%")
    table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def direct_ghg_emission_3_3_1(data):
    # 首先過濾出 scope 為 1 的資料，並依照 scopeName 分組
    scope1_data = {}
    for asset in data:
        for item in asset['companyAsset']:
            if item['scope'] == '1':
                scopeName = item['scopeName']
                if scopeName not in scope1_data:
                    scope1_data[scopeName] = []
                scope1_data[scopeName].extend(item['assets'])

    # 再依照 activityName 分類，每個 activityName 只取一筆
    pre_result = {}
    for scopeName, assets in scope1_data.items():
        pre_result[scopeName] = {}
        seen_activities = set()
        
        for asset in assets:
            activity_name = asset.get('activityName', '')  # 有些資料可能沒有 activityName
            if activity_name and activity_name not in seen_activities:
                seen_activities.add(activity_name)
                pre_result[scopeName][activity_name] = asset.get('fuelType')
    
    result = []
    for scopeName, activities in pre_result.items():
        # 跳過空的活動設施
        if not activities:
            continue
            
        for activityName, fuelType in activities.items():
            # 如果 fuelType 是 None，轉換成空字串
            result.append((
                scopeName,
                activityName,
                "" if fuelType is None else fuelType
            ))
                
    return result



def indirect_ghg_emission_3_3_3(data):
    # 首先過濾出 scope 為 3或4 的資料，並依照 scopeName 分組
    scope_data = {}
    for asset in data:
        for item in asset['companyAsset']:
            if item['scope'] == '3' or item['scope'] == '4':
                scopeName = item['scopeName']
                if scopeName not in scope_data:
                    scope_data[scopeName] = []
                scope_data[scopeName].extend(item['assets'])

    # 再依照 activityName 分類，每個 activityName 只取一筆
    pre_result = {}
    for scopeName, assets in scope_data.items():
        pre_result[scopeName] = {}
        seen_activities = set()
        
        for asset in assets:
            activity_name = asset.get('activityName', '')  # 有些資料可能沒有 activityName
            if activity_name and activity_name not in seen_activities:
                seen_activities.add(activity_name)
                pre_result[scopeName][activity_name] = asset.get('fuelType')
    
    result = []
    for scopeName, activities in pre_result.items():
        # 跳過空的活動設施
        if not activities:
            continue
            
        for activityName, fuelType in activities.items():
            # 如果 fuelType 是 None，轉換成空字串
            result.append((
                scopeName,
                activityName,
                "" if fuelType is None else fuelType
            ))
                
    return result