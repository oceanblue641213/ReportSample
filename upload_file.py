from docx import Document

def replace_text_in_word(file_path, replacements):
    # 打開文檔
    doc = Document(file_path)
    
    dic = {"一": "台北", "二": "桃園", "三": "新竹", "四": "苗栗", "五": "台中"}
    
    first_table = doc.tables[0]
    
    for row in first_table.rows:
        if row._index > 0:
            row._element.getparent().remove(row._element)
    
    for row in dic.keys():
        first_table.add_row()
        
        # for cell in row.cells:
        #     print(f"文字是:{cell.text}，row是:{row._index}")
    
    # 遍歷段落
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                # 获取段落的样式
                paragraph_style = paragraph.style
                
                # 在段落的所有run中查找并替换
                for run in paragraph.runs:
                    if old in run.text:
                        # 替换文字并保留样式
                        run.text = run.text.replace(old, new)
                    
    
    # 遍歷表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)
    
    # 保存新文檔
    doc.save('updated_document.docx')